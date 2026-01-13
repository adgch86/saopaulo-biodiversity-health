from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Titulo
title = doc.add_heading('Resumen del Analisis de Datos - Sao Paulo', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitulo
subtitle = doc.add_paragraph('Fecha: Enero 2026')
subtitle.add_run('\nFuentes: AdaptaBrasil MCTI, DATASUS, IBGE')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ===== SECCION 1 =====
doc.add_heading('1. Descripcion de los Datos', level=1)
doc.add_heading('Variables por Categoria (33 total)', level=2)

table1 = doc.add_table(rows=8, cols=3)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
hdr[0].text = 'Categoria'
hdr[1].text = 'Variables'
hdr[2].text = 'Descripcion'

data1 = [
    ('Riesgos Climaticos', '4', 'flooding_exposure, flooding_risks, hydric_stress_exp/risk'),
    ('Salud', '11', 'persist_* e incidence_* para malaria, dengue, diarrea, leptospirosis, chagas'),
    ('Capacidad Adaptativa (UAI)', '6', 'UAI_housing, UAI_env, UAI_food, UAI_mob, UAI_Crisk'),
    ('Vulnerabilidad Social', '6', 'population, population_preta/indigena/branca, pct_rural'),
    ('Biodiversidad', '3', 'mean/max_species_richness, Vert_rich_risk'),
    ('Seguridad Alimentaria', '1', 'pol_deficit'),
    ('Identificadores', '2', 'COD, Municipality'),
]
for i, (cat, var, desc) in enumerate(data1):
    row = table1.rows[i+1].cells
    row[0].text = cat
    row[1].text = var
    row[2].text = desc

doc.add_paragraph()

# ===== SECCION 2 =====
doc.add_heading('2. Estadisticas Descriptivas', level=1)

doc.add_heading('Riesgos Climaticos (645 municipios)', level=2)
p = doc.add_paragraph()
p.add_run('Exposicion a inundaciones: ').bold = True
p.add_run('Media 0.52 (rango 0.07-0.94)\n')
p.add_run('Riesgo de inundacion: ').bold = True
p.add_run('Media 0.08 (75% de municipios = 0)\n')
p.add_run('Estres hidrico: ').bold = True
p.add_run('Media 0.49 (rango 0-0.77)')

doc.add_heading('Biodiversidad', level=2)
p = doc.add_paragraph()
p.add_run('Riqueza de especies: ').bold = True
p.add_run('Media 623 (rango 252-757)\n')
p.add_run('Riesgo para vertebrados: ').bold = True
p.add_run('Media 9.8 (rango 0-44)')

doc.add_heading('Indice de Acceso Universal (UAI)', level=2)
table2 = doc.add_table(rows=6, cols=4)
table2.style = 'Table Grid'
hdr = table2.rows[0].cells
hdr[0].text = 'Componente'
hdr[1].text = 'Media'
hdr[2].text = 'Min'
hdr[3].text = 'Max'

uai_data = [
    ('UAI_housing', '0.34', '0', '1'),
    ('UAI_env', '0.49', '0', '1'),
    ('UAI_food', '0.41', '0', '1'),
    ('UAI_mob', '0.41', '0', '1'),
    ('UAI_Crisk', '0.32', '0', '1'),
]
for i, row_data in enumerate(uai_data):
    row = table2.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

# ===== SECCION 3 =====
doc.add_heading('3. Correlaciones Clave', level=1)

doc.add_heading('En 645 municipios', level=2)
table3 = doc.add_table(rows=6, cols=4)
table3.style = 'Table Grid'
hdr = table3.rows[0].cells
hdr[0].text = 'Relacion'
hdr[1].text = 'r'
hdr[2].text = 'p'
hdr[3].text = 'Interpretacion'

corr_data = [
    ('Biodiversidad vs Deficit Politico', '-0.567', '<0.001', 'Mayor biodiversidad = menor deficit'),
    ('Riesgo Vertebrados vs Biodiversidad', '-0.763', '<0.001', 'Relacion inversa fuerte'),
    ('UAI_mob vs Poblacion', '+0.725', '<0.001', 'Ciudades grandes mejor movilidad'),
    ('% Rural vs UAI_housing', '-0.415', '<0.001', 'Areas rurales peor acceso'),
    ('Biodiversidad vs Riesgo Inundacion', '+0.272', '<0.001', 'Areas biodiversas mas expuestas'),
]
for i, row_data in enumerate(corr_data):
    row = table3.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

doc.add_heading('En 187 municipios (con datos de salud)', level=2)
table4 = doc.add_table(rows=5, cols=4)
table4.style = 'Table Grid'
hdr = table4.rows[0].cells
hdr[0].text = 'Relacion'
hdr[1].text = 'r'
hdr[2].text = 'p'
hdr[3].text = 'Interpretacion'

health_corr = [
    ('Malaria vs UAI', '-0.593', '<0.001', 'UAI protege contra malaria'),
    ('Diarrea vs UAI', '-0.460', '<0.001', 'UAI protege contra diarrea'),
    ('Dengue vs UAI', '-0.432', '<0.001', 'UAI protege contra dengue'),
    ('Malaria vs % Rural', '+0.600', '<0.001', 'Mayor incidencia en areas rurales'),
]
for i, row_data in enumerate(health_corr):
    row = table4.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

# ===== SECCION 4 =====
doc.add_heading('4. Clasificacion por Cuadrantes', level=1)

table5 = doc.add_table(rows=5, cols=7)
table5.style = 'Table Grid'
hdr = table5.rows[0].cells
headers = ['Cuadrante', 'N', '%', 'UAI', 'Biodiv.', 'Deficit', 'Descripcion']
for i, h in enumerate(headers):
    hdr[i].text = h

cuad_data = [
    ('Q1 - Modelo', '195', '30.2%', '0.56', '690', '0.81', 'Mejor desempeno'),
    ('Q2 - Conservar', '128', '19.8%', '0.24', '680', '0.86', 'Alta biodiv, bajo acceso'),
    ('Q3 - Vulnerable', '192', '29.8%', '0.24', '566', '0.95', 'Intervencion urgente'),
    ('Q4 - Desarrollo', '130', '20.2%', '0.52', '552', '0.89', 'Desarrollo sin conserv.'),
]
for i, row_data in enumerate(cuad_data):
    row = table5.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

doc.add_heading('Municipios Destacados', level=2)
p = doc.add_paragraph()
p.add_run('Q1 (Modelo): ').bold = True
p.add_run('Guarulhos, Campinas, Sao Bernardo do Campo\n')
p.add_run('Q3 (Vulnerable): ').bold = True
p.add_run('Sao Joaquim da Barra, Jardinopolis, Pontal\n')
p.add_run('Q4 (Desarrollo): ').bold = True
p.add_run('Sao Paulo, Santo Andre, Osasco')

doc.add_paragraph()

# ===== SECCION 5 =====
doc.add_heading('5. Analisis de Justicia Ambiental', level=1)

doc.add_heading('Composicion Demografica por Cuadrante', level=2)
table6 = doc.add_table(rows=5, cols=4)
table6.style = 'Table Grid'
hdr = table6.rows[0].cells
hdr[0].text = 'Cuadrante'
hdr[1].text = '% Negra'
hdr[2].text = '% Indigena'
hdr[3].text = '% Rural'

demo_data = [
    ('Q1 - Modelo', '5.35%', '0.11%', '12.5%'),
    ('Q2 - Conservar', '4.74%', '0.21%', '19.2%'),
    ('Q3 - Vulnerable', '4.86%', '0.13%', '14.1%'),
    ('Q4 - Desarrollo', '5.25%', '0.08%', '8.2%'),
]
for i, row_data in enumerate(demo_data):
    row = table6.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

doc.add_heading('Paradoja Biodiversidad-Vulnerabilidad', level=2)
p = doc.add_paragraph()
p.add_run('127 municipios ').bold = True
p.add_run('tienen alta biodiversidad pero bajo acceso a servicios\n')
p.add_run('Deficit politico promedio: 0.86 (alto)\n')
p.add_run('Concentran poblacion rural (19.3%)')

doc.add_paragraph()

# ===== SECCION 6 =====
doc.add_heading('6. Hallazgos Principales', level=1)

p = doc.add_paragraph()
p.add_run('1. El UAI es factor protector: ').bold = True
p.add_run('Correlaciones negativas fuertes con malaria (-0.59), diarrea (-0.46) y dengue (-0.43)\n\n')

p.add_run('2. Biodiversidad asociada a menor deficit politico: ').bold = True
p.add_run('r=-0.57, sugiere que la conservacion esta vinculada a mejor gobernanza\n\n')

p.add_run('3. 192 municipios vulnerables (30%) ').bold = True
p.add_run('requieren intervencion prioritaria - tienen el mayor deficit politico (0.95) y menor biodiversidad\n\n')

p.add_run('4. Areas rurales mas expuestas: ').bold = True
p.add_run('Mayor incidencia de malaria y menor acceso a vivienda\n\n')

p.add_run('5. Poblaciones indigenas ').bold = True
p.add_run('estan significativamente asociadas a areas de mayor biodiversidad')

doc.add_paragraph()

# ===== SECCION 7 =====
doc.add_heading('7. Recomendaciones de Politicas', level=1)

doc.add_heading('Prioridad Alta (Q3 - Vulnerable)', level=2)
p = doc.add_paragraph()
p.add_run('- Inversion en infraestructura basica y UAI\n')
p.add_run('- Fortalecimiento institucional (reducir deficit politico)\n')
p.add_run('- Restauracion ecologica')

doc.add_heading('Prioridad Media (Q2 - Conservar)', level=2)
p = doc.add_paragraph()
p.add_run('- Proteger biodiversidad existente\n')
p.add_run('- Mejorar acceso a servicios sin degradar ecosistemas\n')
p.add_run('- Desarrollo sostenible')

doc.add_heading('Monitoreo (Q1 - Modelo)', level=2)
p = doc.add_paragraph()
p.add_run('- Documentar y replicar buenas practicas\n')
p.add_run('- Mantener equilibrio biodiversidad-desarrollo')

doc.add_paragraph()
doc.add_paragraph('_' * 50)
p = doc.add_paragraph()
p.add_run('Analisis realizado con datos de AdaptaBrasil MCTI, DATASUS e IBGE').italic = True

# Guardar
doc.save('docs/RESUMEN_ANALISIS_METADATA.docx')
print('Documento Word creado: docs/RESUMEN_ANALISIS_METADATA.docx')
