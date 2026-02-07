"""
Create Branco Weiss Proposal Word Document
==========================================
Generates a 5-page max Word document with 2 figures and concise legends.

Author: Science Team - Dr. Adrian David Gonzalez Chaves
Date: January 2026
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
import os

# Base paths
BASE_DIR = r"C:\Users\arlex\Documents\Adrian David"
BRANCO_DIR = os.path.join(BASE_DIR, "admin", "becas", "branco_weiss", "2025_01_14_branco")
OUTPUTS_DIR = os.path.join(BASE_DIR, "outputs", "figures")


def crop_climate_health_figure():
    """Crop the top-right quadrant (Clima vs Salud) from nexo_v5.png"""
    img_path = os.path.join(OUTPUTS_DIR, "nexo_clima_biodiv_salud_v5.png")
    output_path = os.path.join(BRANCO_DIR, "Figure2_Climate_Health.png")

    img = Image.open(img_path)
    # Image is 2081x1754, top-right quadrant
    # Adjusted coordinates to exclude heatmap colorbar artifacts
    # left=1120 to skip the colorbar values from heatmap
    cropped = img.crop((1120, 0, 2081, 877))
    cropped.save(output_path, dpi=(300, 300))
    print(f"Figure 2 cropped and saved to: {output_path}")
    return output_path


def create_document():
    """Create and configure the Word document"""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Configure default paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    return doc


def add_title_and_abstract(doc):
    """Add title and abstract"""
    # Title
    title = doc.add_heading(
        'Resilient Landscapes: Integrating Planetary Health, Nexus Thinking, '
        'and Environmental Justice into Climate Governance for Tropical Cities',
        level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author
    author = doc.add_paragraph()
    author_run = author.add_run('Adrian David Gonzalez Chaves')
    author_run.bold = True
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract heading
    doc.add_heading('Abstract', level=1)

    # Abstract text
    abstract_text = (
        "This project develops a justice-oriented framework for climate adaptation governance "
        "in tropical Latin America by bridging Evidence-Based Policy Making (EBPM), Planetary Health, "
        "and Nexus thinking. Using Sao Paulo State, Brazil (645 municipalities) as the primary laboratory, "
        "with comparative cases in Mexico and Colombia, this research examines relationships among "
        "governance capacity, ecosystem services, health outcomes, and social vulnerability. "
        "Preliminary findings reveal a biodiversity-vulnerability paradox and demonstrate "
        "the value of integrated cross-sectoral analysis for equitable climate adaptation."
    )
    doc.add_paragraph(abstract_text)


def add_research_context(doc):
    """Add Research Context section"""
    doc.add_heading('1. Research Context', level=1)

    p1 = doc.add_paragraph(
        "Climate change poses existential challenges for tropical developing countries, "
        "yet available policy tools remain inadequate. Three critical gaps characterize current research:"
    )

    # Gap 1
    p2 = doc.add_paragraph()
    p2.add_run('Disciplinary fragmentation: ').bold = True
    p2.add_run(
        "Climate adaptation, biodiversity conservation, public health, and social policy "
        "operate in separate silos. Municipal governments receive contradictory guidance "
        "with no framework for integration."
    )

    # Gap 2
    p3 = doc.add_paragraph()
    p3.add_run('Northern bias in EBPM: ').bold = True
    p3.add_run(
        "Evidence-Based Policy Making emerged from OECD contexts assuming robust data "
        "infrastructure rarely present in the Global South."
    )

    # Gap 3
    p4 = doc.add_paragraph()
    p4.add_run('Environmental justice blind spots: ').bold = True
    p4.add_run(
        "Ecosystem service research often ignores distributional questions - "
        "who benefits and who bears risks."
    )


def add_research_aims(doc):
    """Add Research Aims section"""
    doc.add_heading('2. Research Aims', level=1)

    aims = [
        ("Aim 1", "Characterize relationships among governance capacity, ecosystem services, "
                  "health outcomes, and social vulnerability stratified by demographic characteristics."),
        ("Aim 2", "Identify policy gaps and environmental injustices through integrated Nexus analysis."),
        ("Aim 3", "Co-develop decision-support tools enabling visualization across multiple dimensions."),
        ("Aim 4", "Generate transferable methodology for equity-integrated EBPM in tropical developing "
                  "country contexts.")
    ]

    for aim_title, aim_text in aims:
        p = doc.add_paragraph()
        p.add_run(f'{aim_title}: ').bold = True
        p.add_run(aim_text)


def add_methodology(doc):
    """Add Methodological Framework section"""
    doc.add_heading('3. Methodological Framework', level=1)

    doc.add_paragraph(
        "The project employs a mixed-methods approach integrating quantitative analysis of "
        "multi-sectoral databases with qualitative policy analysis:"
    )

    # Urban Adaptability Index
    p = doc.add_paragraph()
    p.add_run('Urban Adaptability Index (UAI): ').bold = True
    p.add_run(
        "Integrates 26 policy indicators across five dimensions: housing, food systems, "
        "environmental management, urban mobility, and climate risk response."
    )

    # Demographic Integration
    p = doc.add_paragraph()
    p.add_run('Demographic Integration: ').bold = True
    p.add_run(
        "Systematic analysis stratified by racial composition and rural/urban distribution "
        "to identify environmental justice patterns."
    )

    # Nexus Analysis
    p = doc.add_paragraph()
    p.add_run('Nexus Analysis: ').bold = True
    p.add_run(
        "Cross-sectoral correlation analysis linking biodiversity, climate risk, disease burden, "
        "governance capacity, and social vulnerability indicators."
    )


def add_figure1(doc):
    """Add Figure 1 with concise legend"""
    fig1_path = os.path.join(BRANCO_DIR, "Figure1_DataIntegration.png")

    # Add image
    doc.add_picture(fig1_path, width=Inches(6.0))

    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add legend (concise)
    legend = doc.add_paragraph()
    legend.add_run('Figure 1. ').bold = True
    legend.add_run(
        'Integrated analysis across 645 municipalities of Sao Paulo State. '
        '(A) Biodiversity-vulnerability paradox (r=0.23, p<0.001). '
        '(B) Dilution effect: biodiversity reduces disease incidence (r=-0.43, p<0.001). '
        '(C) Governance gap in vulnerable municipalities.'
    )
    legend.paragraph_format.space_after = Pt(12)


def add_preliminary_results(doc):
    """Add Preliminary Results section"""
    doc.add_heading('4. Preliminary Results', level=1)

    doc.add_paragraph(
        "Analysis of Sao Paulo State data reveals patterns that emerge only through "
        "integrated cross-sectoral analysis (Figure 1):"
    )

    # Result 1
    p = doc.add_paragraph()
    p.add_run('Biodiversity-Vulnerability Paradox: ').bold = True
    p.add_run(
        "Municipalities with higher social vulnerability tend to have greater species richness "
        "(r=0.23, p<0.001), suggesting conservation-development trade-offs requiring policy attention."
    )

    # Result 2
    p = doc.add_paragraph()
    p.add_run('Dilution Effect: ').bold = True
    p.add_run(
        "Higher biodiversity correlates with lower vector-borne disease incidence "
        "(dengue: r=-0.43, diarrhea: r=-0.45; p<0.001), demonstrating ecosystem-health linkages "
        "invisible to siloed analyses."
    )

    # Result 3
    p = doc.add_paragraph()
    p.add_run('Climate-Health Nexus: ').bold = True
    p.add_run(
        "Higher climate risk associates with greater disease burden (r=-0.54, p<0.001), "
        "with biodiversity acting as a potential moderating factor (Figure 2)."
    )


def add_figure2(doc, fig2_path):
    """Add Figure 2 with concise legend"""
    # Add image (smaller than Figure 1)
    doc.add_picture(fig2_path, width=Inches(3.5))

    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add legend (very concise)
    legend = doc.add_paragraph()
    legend.add_run('Figure 2. ').bold = True
    legend.add_run(
        'Climate-health nexus: higher climate risk correlates with greater disease burden '
        '(r=-0.54, p<0.001). Color indicates biodiversity index, illustrating the cross-sectoral '
        'integration central to this project.'
    )
    legend.paragraph_format.space_after = Pt(12)


def add_significance(doc):
    """Add Significance and Innovation section"""
    doc.add_heading('5. Significance and Innovation', level=1)

    doc.add_paragraph(
        "This project advances climate adaptation science through three innovations:"
    )

    innovations = [
        ("Integrated Framework", "Bridges Planetary Health, Nexus thinking, and Environmental Justice "
                                 "into a coherent analytical approach for tropical contexts."),
        ("South-South Methodology", "Develops EBPM tools adapted to data-scarce environments, "
                                    "transferable across Latin America and beyond."),
        ("Justice-Centered Analysis", "Explicitly incorporates demographic stratification to identify "
                                      "and address environmental inequities.")
    ]

    for title, text in innovations:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(text)


def add_timeline(doc):
    """Add Timeline section"""
    doc.add_heading('6. Timeline', level=1)

    timeline_items = [
        ("Year 1", "Complete data integration for Sao Paulo; develop analytical framework"),
        ("Year 2", "Extend analysis to Mexico and Colombia; validate cross-context patterns"),
        ("Year 3", "Develop decision-support platform; policy co-design workshops"),
        ("Year 4", "Dissemination, training modules, and scaling strategies"),
        ("Year 5", "Policy impact assessment and methodology refinement")
    ]

    for year, activities in timeline_items:
        p = doc.add_paragraph()
        p.add_run(f'{year}: ').bold = True
        p.add_run(activities)


def add_references(doc):
    """Add References section"""
    doc.add_heading('References', level=1)

    refs = [
        "Barreto, J.R., et al. (2025). Indigenous Territories can safeguard human health depending on the landscape structure and legal status. Communications Earth & Environment, 6:719.",
        "Cairney, P., & Oliver, K. (2017). Evidence-based policymaking is not like evidence-based medicine. Health Res Policy Syst, 15(1), 35.",
        "Levers, C., et al. (2025). Different places, different challenges: mapping global variations in agrifood-system burdens. Environmental Research Letters, 20:124051.",
        "Neder, E.A., et al. (2021). Urban adaptation index: assessing cities readiness for climate change. Clim Change, 166, 16.",
        "Whitmee, S., et al. (2015). Safeguarding human health in the Anthropocene. The Lancet, 386, 1973-2028."
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)


def main():
    """Main execution function"""
    print("=" * 60)
    print("Creating Branco Weiss Proposal Word Document")
    print("=" * 60)

    # Step 1: Crop Figure 2
    print("\nStep 1: Cropping Figure 2 from nexo image...")
    fig2_path = crop_climate_health_figure()

    # Step 2: Create document
    print("\nStep 2: Creating Word document...")
    doc = create_document()

    # Step 3: Add content sections
    print("\nStep 3: Adding content sections...")
    add_title_and_abstract(doc)
    add_research_context(doc)
    add_research_aims(doc)
    add_methodology(doc)
    add_figure1(doc)
    add_preliminary_results(doc)
    add_figure2(doc, fig2_path)
    add_significance(doc)
    add_timeline(doc)
    add_references(doc)

    # Step 4: Save document
    output_path = os.path.join(BRANCO_DIR, "BrancoWeiss_Proposal.docx")
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    print("\nDone! Please review the document and verify it is <= 5 pages.")


if __name__ == "__main__":
    main()
