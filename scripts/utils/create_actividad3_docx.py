"""
Genera documento Word de la Actividad 3 actualizada
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilos
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# TITULO
title = doc.add_heading('Workshop SEMIL-USP: Actividad 3', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Nexus Assessment and Overarching Results')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(14)

doc.add_paragraph('Fecha de actualizacion: 2026-02-03')
doc.add_paragraph('Basado en: Analisis H1-H6 (Framework de Predictores de Gobernanza)')
doc.add_paragraph()

# RESUMEN EJECUTIVO
doc.add_heading('Resumen Ejecutivo', level=1)
doc.add_paragraph('Esta actividad presenta los resultados integrados del nexus assessment con 104+ variables para los 645 municipios de Sao Paulo.')

p = doc.add_paragraph()
p.add_run('IMPORTANTE: ').bold = True
p.add_run('Los hallazgos difieren significativamente del analisis original debido a:')

doc.add_paragraph('1. Validacion rigurosa del efecto de confusion por urbanizacion', style='List Number')
doc.add_paragraph('2. Inversion de la logica causal (riesgos genera gobernanza)', style='List Number')
doc.add_paragraph('3. Identificacion de que dengue NO responde a cobertura forestal', style='List Number')

# ESTRUCTURA
doc.add_heading('Estructura de la Actividad', level=1)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Componente'
hdr[1].text = 'Duracion'
hdr[2].text = 'Formato'

data = [
    ('Presentacion de resultados', '15 min', 'Plenaria'),
    ('Discusion en grupos', '25 min', '4-5 grupos'),
    ('Sintesis y debate', '20 min', 'Plenaria integrada'),
    ('TOTAL', '60 min', ''),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

# PARTE 1
doc.add_heading('PARTE 1: Presentacion de Resultados (15 min)', level=1)

# Hallazgo 1
doc.add_heading('Hallazgo 1: Dengue es URBANO, no forestal', level=2)
p = doc.add_paragraph()
p.add_run('Pregunta original: ').bold = True
p.add_run('La biodiversidad protege contra dengue?')

p = doc.add_paragraph()
p.add_run('Respuesta: NO. ').bold = True
p.add_run('El efecto aparente esta CONFUNDIDO por urbanizacion.')

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Analisis'
table.rows[0].cells[1].text = 'Resultado'
table.rows[1].cells[0].text = 'Correlacion simple bosque-dengue'
table.rows[1].cells[1].text = 'r = -0.45 (parecia protector)'
table.rows[2].cells[0].text = 'Controlando por poblacion y ruralidad'
table.rows[2].cells[1].text = 'Beta cambia 36%, deja de ser significativo'
table.rows[3].cells[0].text = 'Mejor predictor de dengue'
table.rows[3].cells[1].text = '% Rural (Beta = -0.225, R2m = 43%)'

p = doc.add_paragraph()
p.add_run('Implicacion: ').bold = True
p.add_run('Las intervenciones contra dengue deben focalizarse en areas urbanas. La conservacion forestal NO es una estrategia efectiva contra dengue.')

doc.add_paragraph()

# Hallazgo 2
doc.add_heading('Hallazgo 2: Malaria SI responde a bosque (pero al reves)', level=2)
p = doc.add_paragraph()
p.add_run('Respuesta: ').bold = True
p.add_run('Mas bosque = MAS malaria. Este efecto es ROBUSTO (no cambia al controlar por poblacion).')

p = doc.add_paragraph()
p.add_run('Explicacion ecologica: ').bold = True
p.add_run('El vector de malaria (Anopheles) requiere ecosistemas forestales. Trade-off real entre conservacion y salud.')

doc.add_paragraph()

# Hallazgo 3
doc.add_heading('Hallazgo 3: La Paradoja de Gobernanza', level=2)
p = doc.add_paragraph()
p.add_run('Pregunta: ').bold = True
p.add_run('La gobernanza reduce los riesgos climaticos?')

p = doc.add_paragraph()
p.add_run('Respuesta: Al contrario. ').bold = True
p.add_run('Municipios con mayor riesgo tienen MAYOR gobernanza.')

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Relacion'
table.rows[0].cells[1].text = 'Coeficiente'
table.rows[0].cells[2].text = 'Interpretacion'
table.rows[1].cells[0].text = 'Riesgo Fuego -> Gobernanza'
table.rows[1].cells[1].text = 'Beta = +0.315'
table.rows[1].cells[2].text = 'Mas riesgo = mas gobernanza'
table.rows[2].cells[0].text = 'Riesgo Inundacion -> UAI Clima'
table.rows[2].cells[1].text = 'Beta = +1.033'
table.rows[2].cells[2].text = 'Inversion reactiva'
table.rows[3].cells[0].text = 'Riesgo Climatico -> UAI Movilidad'
table.rows[3].cells[1].text = 'Beta = +1.120'
table.rows[3].cells[2].text = 'Desarrollo post-crisis'

p = doc.add_paragraph()
p.add_run('Interpretacion: ').bold = True
p.add_run('La gobernanza es REACTIVA, no preventiva.')

doc.add_paragraph()

# Hallazgo 4
doc.add_heading('Hallazgo 4: La Pobreza Domina Todo', level=2)
p = doc.add_paragraph()
p.add_run('La pobreza es el predictor mas fuerte de gobernanza (explica 27% de la varianza).')

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Variable Respuesta'
table.rows[0].cells[1].text = 'Mejor Predictor'
table.rows[0].cells[2].text = 'R2m'
table.rows[1].cells[0].text = 'Gobernanza general'
table.rows[1].cells[1].text = '% Pobreza'
table.rows[1].cells[2].text = '27%'
table.rows[2].cells[0].text = 'UAI Vivienda'
table.rows[2].cells[1].text = '% Pobreza'
table.rows[2].cells[2].text = '21%'
table.rows[3].cells[0].text = 'UAI Movilidad'
table.rows[3].cells[1].text = '% Pobreza'
table.rows[3].cells[2].text = '19%'
table.rows[4].cells[0].text = 'Dengue'
table.rows[4].cells[1].text = '% Rural'
table.rows[4].cells[2].text = '43%'
table.rows[5].cells[0].text = 'Mort. Cardiovascular'
table.rows[5].cells[1].text = '% Pob. Negra'
table.rows[5].cells[2].text = '25%'

p = doc.add_paragraph()
p.add_run('Implicacion: ').bold = True
p.add_run('Sin reducir pobreza, las intervenciones ambientales tienen efecto limitado.')

doc.add_paragraph()

# Hallazgo 5
doc.add_heading('Hallazgo 5: Lo que SI funciona', level=2)
doc.add_paragraph('Servicios ecosistemicos con efectos robustos:')

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Relacion'
table.rows[0].cells[1].text = 'Coeficiente'
table.rows[0].cells[2].text = 'Validacion'
table.rows[1].cells[0].text = 'Bosque -> reduce deficit polinizacion'
table.rows[1].cells[1].text = 'r = -0.77'
table.rows[1].cells[2].text = 'Muy fuerte'
table.rows[2].cells[0].text = 'Bosque -> reduce hosp. respiratorias'
table.rows[2].cells[1].text = 'Beta = -0.145'
table.rows[2].cells[2].text = 'Robusto'
table.rows[3].cells[0].text = 'Bosque -> reduce diarrea'
table.rows[3].cells[1].text = 'Beta = -0.342'
table.rows[3].cells[2].text = 'Robusto'
table.rows[4].cells[0].text = 'Bosque -> dengue'
table.rows[4].cells[1].text = 'No significativo'
table.rows[4].cells[2].text = 'Confundido'

doc.add_paragraph()

# Cuadrantes
doc.add_heading('Clasificacion de Municipios - 4 Cuadrantes', level=2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Cuadrante'
table.rows[0].cells[1].text = 'N municipios'
table.rows[0].cells[2].text = '% Poblacion'
table.rows[0].cells[3].text = 'Caracteristica'
table.rows[1].cells[0].text = 'Q1 - Modelo'
table.rows[1].cells[1].text = '~160'
table.rows[1].cells[2].text = '~45%'
table.rows[1].cells[3].text = 'Alta gob + Baja vuln'
table.rows[2].cells[0].text = 'Q2 - Riesgo'
table.rows[2].cells[1].text = '~160'
table.rows[2].cells[2].text = '~20%'
table.rows[2].cells[3].text = 'Alta gob + Alta vuln'
table.rows[3].cells[0].text = 'Q3 - Critico'
table.rows[3].cells[1].text = '~160'
table.rows[3].cells[2].text = '~15%'
table.rows[3].cells[3].text = 'Baja gob + Alta vuln'
table.rows[4].cells[0].text = 'Q4 - Potencial'
table.rows[4].cells[1].text = '~165'
table.rows[4].cells[2].text = '~20%'
table.rows[4].cells[3].text = 'Baja gob + Baja vuln'

doc.add_paragraph()

# PARTE 2
doc.add_heading('PARTE 2: Discusion en Grupos (25 min)', level=1)
doc.add_paragraph('Dividir participantes en 4-5 grupos de 4-5 personas. Cada grupo recibe 2 preguntas provocadoras.')

doc.add_heading('Preguntas Provocadoras', level=2)

p = doc.add_paragraph()
p.add_run('Grupo 1: Dengue y Urbanizacion').bold = True
doc.add_paragraph('1. Si dengue es fundamentalmente un problema urbano, tiene sentido incluir conservacion de biodiversidad como estrategia de control de vectores?')
doc.add_paragraph('2. Como deberian diferir las estrategias de salud publica entre municipios urbanos vs rurales?')

p = doc.add_paragraph()
p.add_run('Grupo 2: La Paradoja de Gobernanza').bold = True
doc.add_paragraph('3. Sabiendo que la gobernanza es reactiva, como podriamos disenar incentivos para accion preventiva?')
doc.add_paragraph('4. Es posible o deseable que la gobernanza sea proactiva? Que barreras institucionales lo impiden?')

p = doc.add_paragraph()
p.add_run('Grupo 3: Pobreza como Factor Dominante').bold = True
doc.add_paragraph('5. Si la pobreza domina todas las relaciones, deberiamos enfocarnos primero en reducir pobreza antes de intervenciones ambientales?')
doc.add_paragraph('6. Como evitar que las intervenciones de conservacion aumenten las desigualdades existentes?')

p = doc.add_paragraph()
p.add_run('Grupo 4: Trade-offs Biodiversidad-Salud').bold = True
doc.add_paragraph('7. La malaria aumenta con la cobertura forestal - existen trade-offs inevitables entre conservacion y salud?')
doc.add_paragraph('8. Como comunicar al publico que mas naturaleza no siempre significa mejor salud?')

p = doc.add_paragraph()
p.add_run('Grupo 5: Implicaciones para Politicas').bold = True
doc.add_paragraph('9. Que recomendaciones especificas harian para el Plan Estadual de Adaptacion Climatica de Sao Paulo?')
doc.add_paragraph('10. Como integrar estos hallazgos en la planificacion municipal sin recursos tecnicos especializados?')

doc.add_paragraph()

# PARTE 3
doc.add_heading('PARTE 3: Sintesis y Debate (20 min)', level=1)

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Tiempo'
table.rows[0].cells[1].text = 'Actividad'
table.rows[1].cells[0].text = '0-12 min'
table.rows[1].cells[1].text = 'Cada grupo presenta sus 2-3 puntos clave'
table.rows[2].cells[0].text = '12-18 min'
table.rows[2].cells[1].text = 'Debate abierto'
table.rows[3].cells[0].text = '18-20 min'
table.rows[3].cells[1].text = 'Sintesis del facilitador'

doc.add_heading('Puntos a Destacar en la Sintesis', level=2)
doc.add_paragraph('1. Complejidad no significa paralisis - los resultados ofrecen direcciones claras.')
doc.add_paragraph('2. Especificidad importa - cada enfermedad y municipio requiere estrategias diferenciadas.')
doc.add_paragraph('3. La pobreza es el elefante en la habitacion - sin abordarla, las intervenciones tienen efecto limitado.')
doc.add_paragraph('4. Gobernanza reactiva es una oportunidad - podemos disenar sistemas que anticipen.')
doc.add_paragraph('5. Trade-offs son reales - debemos ser honestos y disenar compensaciones justas.')

# ANEXO
doc.add_page_break()
doc.add_heading('ANEXO: Comparacion Original vs Actualizado', level=1)

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Tema'
table.rows[0].cells[1].text = 'Original'
table.rows[0].cells[2].text = 'Actualizado'
table.rows[1].cells[0].text = 'Efecto dilucion dengue'
table.rows[1].cells[1].text = 'Confirmado (r=-0.45)'
table.rows[1].cells[2].text = 'CONFUNDIDO por urbanizacion'
table.rows[2].cells[0].text = 'Mejor predictor dengue'
table.rows[2].cells[1].text = 'Biodiversidad'
table.rows[2].cells[2].text = '% Rural (urbanizacion)'
table.rows[3].cells[0].text = 'Efecto bosque-malaria'
table.rows[3].cells[1].text = 'No mencionado'
table.rows[3].cells[2].text = 'Positivo y robusto'
table.rows[4].cells[0].text = 'Paradoja gobernanza'
table.rows[4].cells[1].text = 'Por urbanizacion'
table.rows[4].cells[2].text = 'Gobernanza es REACTIVA'
table.rows[5].cells[0].text = 'Bosque-polinizacion'
table.rows[5].cells[1].text = 'Mencionado'
table.rows[5].cells[2].text = 'Confirmado (r=-0.77)'
table.rows[6].cells[0].text = 'Factor dominante'
table.rows[6].cells[1].text = 'Vulnerabilidad'
table.rows[6].cells[2].text = 'Pobreza (R2m=27%)'
table.rows[7].cells[0].text = 'Mensaje conservacion'
table.rows[7].cells[1].text = 'Protege salud'
table.rows[7].cells[2].text = 'Depende de enfermedad'

doc.add_paragraph()
doc.add_paragraph('Documento preparado por Science Team para Workshop SEMIL-USP - Febrero 2026')

# Guardar
output_path = r'C:\Users\arlex\Documents\Adrian David\docs\WORKSHOP_ACTIVIDAD_3_ACTUALIZADA.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
